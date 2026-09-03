# -*- coding: utf-8 -*-
"""
Full Generator for Chapter 2: ทฤษฎีและงานวิจัยที่เกี่ยวข้อง
Outputs:
1. 01_เล่มรายงานโครงงาน_Thesis/บทที่_2_ทฤษฎีและงานวิจัยที่เกี่ยวข้อง_ฉบับปรับปรุงสมบูรณ์.docx
2. 01_เล่มรายงานโครงงาน_Thesis/บทที่_2_ทฤษฎีและงานวิจัยที่เกี่ยวข้อง_ฉบับปรับปรุงสมบูรณ์.pdf
"""

import os
import sys
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
import latex2mathml.converter as l2m
import lxml.etree as ET
import win32com.client

# Load XSLT
xslt_path = r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL'
xslt = ET.parse(xslt_path)
transform = ET.XSLT(xslt)

def latex_to_omml(latex_str):
    mml = l2m.convert(latex_str)
    dom = ET.fromstring(mml)
    omml = transform(dom)
    xml_bytes = ET.tostring(omml.getroot())
    return parse_xml(xml_bytes)

def set_run_font(run, font_name="TH Sarabun PSK", size_pt=16, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def add_chapter_title(doc, chapter_num, chapter_title):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(6)
    p1.paragraph_format.line_spacing = 1.15
    r1 = p1.add_run(chapter_num)
    set_run_font(r1, size_pt=18, bold=True)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(18)
    p2.paragraph_format.line_spacing = 1.15
    r2 = p2.add_run(chapter_title)
    set_run_font(r2, size_pt=18, bold=True)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size_pt=16, bold=True)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size_pt=16, bold=True)
    return p

def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size_pt=15, bold=True)
    return p

def add_body_p(doc, text, bold_prefix=None, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.THAI_JUSTIFY if hasattr(WD_ALIGN_PARAGRAPH, 'THAI_JUSTIFY') else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_run_font(r_pre, size_pt=16, bold=True)
    
    r = p.add_run(text)
    set_run_font(r, size_pt=16, bold=False)
    return p

def add_bullet_p(doc, bold_prefix, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    
    r_bullet = p.add_run("•  ")
    set_run_font(r_bullet, size_pt=16, bold=True)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_run_font(r_pre, size_pt=16, bold=True)
        
    r = p.add_run(text)
    set_run_font(r, size_pt=16, bold=False)
    return p

def add_numbered_item_p(doc, num_str, bold_prefix, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    
    r_num = p.add_run(num_str + " ")
    set_run_font(r_num, size_pt=16, bold=True)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_run_font(r_pre, size_pt=16, bold=True)
        
    r = p.add_run(text)
    set_run_font(r, size_pt=16, bold=False)
    return p

def add_equation(doc, latex_str, eq_num_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    pPr = p._p.get_or_add_pPr()
    
    tabs_xml = r'''
    <w:tabs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:tab w:val="center" w:pos="4150"/>
      <w:tab w:val="right" w:pos="8300"/>
    </w:tabs>
    '''
    pPr.append(parse_xml(tabs_xml))
    
    r1 = p.add_run()
    r1._r.append(parse_xml(r'<w:tab xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
    
    omml_elem = latex_to_omml(latex_str)
    p._p.append(omml_elem)
    
    r2 = p.add_run()
    r2._r.append(parse_xml(r'<w:tab xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
    
    r3 = p.add_run(eq_num_str)
    set_run_font(r3, size_pt=16, bold=False)
    return p

def add_flow_box(doc, title, flow_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(5.8)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = r'''
    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single" w:sz="12" w:space="0" w:color="2563EB"/>
      <w:left w:val="single" w:sz="12" w:space="0" w:color="2563EB"/>
      <w:bottom w:val="single" w:sz="12" w:space="0" w:color="2563EB"/>
      <w:right w:val="single" w:sz="12" w:space="0" w:color="2563EB"/>
    </w:tcBorders>
    '''
    shd_xml = r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="F8FAFC"/>'
    mar_xml = r'''
    <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:w="160" w:type="dxa"/>
      <w:left w:w="240" w:type="dxa"/>
      <w:bottom w:w="160" w:type="dxa"/>
      <w:right w:w="240" w:type="dxa"/>
    </w:tcMar>
    '''
    tcPr.append(parse_xml(borders_xml))
    tcPr.append(parse_xml(shd_xml))
    tcPr.append(parse_xml(mar_xml))
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"【 {title} 】\n")
    set_run_font(r1, size_pt=15, bold=True, color=RGBColor(37, 99, 235))
    r2 = p.add_run(flow_text)
    set_run_font(r2, size_pt=14, bold=False, color=RGBColor(15, 23, 42))

def style_booktabs_table(table, col_widths, header_bg="F1F5F9"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    borders_xml = r'''
    <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
      <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
      <w:left w:val="none"/>
      <w:right w:val="none"/>
      <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
      <w:insideV w:val="none"/>
    </w:tblBorders>
    '''
    tblPr.append(parse_xml(borders_xml))
    
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
        if r_idx == 0:
            trPr.append(parse_xml(r'<w:tblHeader xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))
            
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]
            tcPr = cell._tc.get_or_add_tcPr()
            
            mar_xml = r'''
            <w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
              <w:top w:w="120" w:type="dxa"/>
              <w:bottom w:w="120" w:type="dxa"/>
              <w:left w:w="140" w:type="dxa"/>
              <w:right w:w="140" w:type="dxa"/>
            </w:tcMar>
            '''
            tcPr.append(parse_xml(mar_xml))
            
            if r_idx == 0:
                shd_xml = f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{header_bg}"/>'
                b_xml = r'''
                <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>
                </w:tcBorders>
                '''
                tcPr.append(parse_xml(shd_xml))
                tcPr.append(parse_xml(b_xml))

print("Helper functions ready.")
